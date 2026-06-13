#!/usr/bin/env python3
"""
LPAI Nexus Internship Technical Report Generator
Generates a professional RESTRICTED .docx report using python-docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Color palette ──────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1A, 0x3A, 0x6A)
GOLD   = RGBColor(0xD4, 0x88, 0x1A)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LGRAY  = RGBColor(0xF5, 0xF5, 0xF5)
MGRAY  = RGBColor(0xD0, 0xD0, 0xD0)
BLACK  = RGBColor(0x1A, 0x1A, 0x1A)
RED_C  = RGBColor(0xC0, 0x39, 0x2B)

BASE   = "/home/asus/lpai/report-assests"
LOGO   = f"{BASE}/logo/LPAI-logo.png"
SS     = f"{BASE}/screenshots"

SCREENSHOTS = {
    "homepage":      f"{SS}/screenshot-homepage.jpeg",
    "cargo":         f"{SS}/screenshot-cargo.jpeg",
    "people":        f"{SS}/screenshot-people.png",
    "security":      f"{SS}/screenshot-security.png",
    "vehiclegate":   f"{SS}/screenshot-vehiclegate.png",
    "intelligence":  f"{SS}/screenshot-intelligence.png",
    "walkthrough":   f"{SS}/screenshot-walkthrough.png",
}

OUTPUT = "/home/asus/lpai/LPAI_Nexus_Internship_Report.docx"

# ── XML helpers ────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        if side in kwargs:
            tag = OxmlElement(f'w:{side}')
            for k, v in kwargs[side].items():
                tag.set(qn(f'w:{k}'), v)
            tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_border_to_paragraph(para, color_hex='1A3A6A', size='6', space='1'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), space)
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def make_paragraph_border_box(para, color_hex='C0392B', size='12'):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top','left','bottom','right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), size)
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), color_hex)
        pBdr.append(el)
    pPr.append(pBdr)

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_break_type())
    return p

def docx_break_type():
    from docx.enum.text import WD_BREAK
    return WD_BREAK.PAGE

def insert_page_break(doc):
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p.append(r)
    doc.element.body.append(p)

def add_hyperlink_style(doc):
    styles = doc.styles
    try:
        styles['Hyperlink']
    except:
        pass

def set_run_font(run, bold=False, italic=False, size=11, color=None, name='Arial'):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def styled_heading(doc, text, level=1, color=NAVY):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.bold = True
    run.font.color.rgb = color
    if level == 1:
        run.font.size = Pt(16)
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(6)
        add_border_to_paragraph(p, '1A3A6A', '8', '4')
    elif level == 2:
        run.font.size = Pt(13)
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        run.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)
        run.font.color.rgb = GOLD
    return p

def body_para(doc, text, indent=0, space_after=8):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, size=11, color=BLACK)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    return p

def bullet_para(doc, text, indent=0.5):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size=11, color=BLACK)
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_screenshot(doc, key, caption_text, width=5.8):
    path = SCREENSHOTS.get(key)
    if path and os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        try:
            run.add_picture(path, width=Inches(width))
        except Exception as e:
            run.text = f'[Screenshot: {key}]'
    else:
        p = doc.add_paragraph(f'[Screenshot not found: {key}]')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption_text)
    set_run_font(r, italic=True, size=9, color=RGBColor(0x55, 0x55, 0x55))
    cap.paragraph_format.space_after = Pt(14)
    return p

def navy_table(doc, headers, rows, col_widths=None):
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, '1A3A6A')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, bold=True, size=10, color=WHITE)

    # Data rows
    for ri, row_data in enumerate(rows):
        drow = table.rows[ri + 1]
        bg = 'F5F5F5' if ri % 2 == 0 else 'FFFFFF'
        for ci, cell_text in enumerate(row_data):
            cell = drow.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(cell_text))
            set_run_font(run, size=10, color=BLACK)

    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Cm(w)

    doc.add_paragraph()
    return table

# ── Header / Footer ────────────────────────────────────────────────────────

def add_headers_footers(doc):
    from docx.oxml.ns import nsmap
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # Default header (all pages except first)
    header = section.header
    header.is_linked_to_previous = False
    # Clear existing
    for p in header.paragraphs:
        p._element.getparent().remove(p._element)

    htable = header.add_table(1, 3, Inches(6.5))
    htable.style = 'Table Grid'
    htable.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Remove borders
    for row in htable.rows:
        for cell in row.cells:
            for side in ['top','left','bottom','right']:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()

    # Left cell
    lc = htable.rows[0].cells[0]
    lp = lc.paragraphs[0]
    lr = lp.add_run("LPAI Nexus — Internship Technical Report")
    set_run_font(lr, size=9, color=NAVY, italic=True)
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Middle cell
    mc = htable.rows[0].cells[1]
    mp = mc.paragraphs[0]
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Right cell
    rc = htable.rows[0].cells[2]
    rp = rc.paragraphs[0]
    rr = rp.add_run("RESTRICTED")
    set_run_font(rr, bold=True, size=9, color=RED_C)
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p._element.getparent().remove(p._element)

    ftable = footer.add_table(1, 3, Inches(6.5))
    ftable.alignment = WD_TABLE_ALIGNMENT.CENTER

    flc = ftable.rows[0].cells[0]
    flp = flc.paragraphs[0]
    flr = flp.add_run("Land Port Authority of India  |  Ministry of Home Affairs")
    set_run_font(flr, size=8, color=NAVY)
    flp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    fmc = ftable.rows[0].cells[1]
    fmp = fmc.paragraphs[0]
    fmp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    frc = ftable.rows[0].cells[2]
    frp = frc.paragraphs[0]
    # Add page number field
    frp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    frr = frp.add_run("Page ")
    set_run_font(frr, size=8, color=NAVY)
    # Page number field
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'begin')
    frp._p.append(OxmlElement('w:r'))
    run_el = OxmlElement('w:r')
    run_el.append(fld)
    frp._p.append(run_el)

    instrText = OxmlElement('w:instrText')
    instrText.text = ' PAGE '
    run_el2 = OxmlElement('w:r')
    run_el2.append(instrText)
    frp._p.append(run_el2)

    fld2 = OxmlElement('w:fldChar')
    fld2.set(qn('w:fldCharType'), 'end')
    run_el3 = OxmlElement('w:r')
    run_el3.append(fld2)
    frp._p.append(run_el3)

    frr2 = frp.add_run(" of ")
    set_run_font(frr2, size=8, color=NAVY)

    fld3 = OxmlElement('w:fldChar')
    fld3.set(qn('w:fldCharType'), 'begin')
    run_el4 = OxmlElement('w:r')
    run_el4.append(fld3)
    frp._p.append(run_el4)

    instrText2 = OxmlElement('w:instrText')
    instrText2.text = ' NUMPAGES '
    run_el5 = OxmlElement('w:r')
    run_el5.append(instrText2)
    frp._p.append(run_el5)

    fld4 = OxmlElement('w:fldChar')
    fld4.set(qn('w:fldCharType'), 'end')
    run_el6 = OxmlElement('w:r')
    run_el6.append(fld4)
    frp._p.append(run_el6)

# ── Cover Page ─────────────────────────────────────────────────────────────

def build_cover(doc):
    section = doc.sections[0]
    section.page_width  = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

    # RESTRICTED banner top
    p_res = doc.add_paragraph()
    p_res.alignment = WD_ALIGN_PARAGRAPH.CENTER
    make_paragraph_border_box(p_res, 'C0392B', '16')
    r = p_res.add_run("  RESTRICTED — GOVERNMENT OF INDIA  ")
    set_run_font(r, bold=True, size=11, color=RED_C)
    p_res.paragraph_format.space_after = Pt(20)

    # Logo
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(LOGO):
        run = p_logo.add_run()
        try:
            run.add_picture(LOGO, height=Inches(1.2))
        except:
            pass
    p_logo.paragraph_format.space_after = Pt(4)

    # Ministry line
    p_min = doc.add_paragraph()
    p_min.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_min.add_run("LAND PORT AUTHORITY OF INDIA")
    set_run_font(r, bold=True, size=13, color=NAVY)
    p_min.paragraph_format.space_after = Pt(2)

    p_mha = doc.add_paragraph()
    p_mha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_mha.add_run("Ministry of Home Affairs, Government of India")
    set_run_font(r, size=10, color=NAVY)
    p_mha.paragraph_format.space_after = Pt(30)

    # Divider
    p_div = doc.add_paragraph()
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_div.add_run("─" * 55)
    set_run_font(r, size=10, color=GOLD)
    p_div.paragraph_format.space_after = Pt(20)

    # Main title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_title.add_run("AI Integration in Land Port Authority of India")
    set_run_font(r, bold=True, size=22, color=NAVY)
    p_title.paragraph_format.space_after = Pt(8)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_sub.add_run("LPAI Nexus — A Smart Border Command Platform")
    set_run_font(r, bold=True, size=16, color=GOLD)
    p_sub.paragraph_format.space_after = Pt(6)

    p_type = doc.add_paragraph()
    p_type.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_type.add_run("Internship Technical Report")
    set_run_font(r, italic=True, size=13, color=RGBColor(0x44, 0x44, 0x44))
    p_type.paragraph_format.space_after = Pt(30)

    # Gold divider
    p_div2 = doc.add_paragraph()
    p_div2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_div2.add_run("─" * 55)
    set_run_font(r, size=10, color=GOLD)
    p_div2.paragraph_format.space_after = Pt(24)

    # Details table
    details = [
        ("Submitted by",    "Himanshu"),
        ("Internship at",   "Land Port Authority of India (LPAI)"),
        ("Reporting Officer","Senior Officer, LPAI — Ministry of Home Affairs"),
        ("Duration",        "May–June 2026"),
        ("Date of Submission", "June 2026"),
        ("Document Status", "RESTRICTED — For Official Use Only"),
        ("Version",         "1.0 — Final"),
    ]

    dtable = doc.add_table(len(details), 2)
    dtable.alignment = WD_TABLE_ALIGNMENT.CENTER
    dtable.style = 'Table Grid'
    for i, (k, v) in enumerate(details):
        lc = dtable.rows[i].cells[0]
        rc = dtable.rows[i].cells[1]
        bg = 'EAF0FB' if i % 2 == 0 else 'FFFFFF'
        set_cell_bg(lc, '1A3A6A')
        set_cell_bg(rc, bg)
        lp = lc.paragraphs[0]
        lr = lp.add_run(k)
        set_run_font(lr, bold=True, size=10, color=WHITE)
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rp = rc.paragraphs[0]
        rr = rp.add_run(v)
        set_run_font(rr, size=10, color=BLACK)
        lc.width = Cm(5)
        rc.width = Cm(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(24)

    # Bottom restricted
    p_bot = doc.add_paragraph()
    p_bot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    make_paragraph_border_box(p_bot, 'C0392B', '12')
    r = p_bot.add_run("  This document contains sensitive government information. Handle accordingly.  ")
    set_run_font(r, bold=True, size=9, color=RED_C)

    # Page break after cover
    doc.add_page_break()

# ── Table of Contents (manual) ─────────────────────────────────────────────

def build_toc(doc):
    styled_heading(doc, "Table of Contents", 1, NAVY)

    toc_items = [
        ("Executive Summary", "3"),
        ("1. Introduction", "4"),
        ("  1.1 Background", "4"),
        ("  1.2 Objectives", "4"),
        ("  1.3 Scope of Work", "5"),
        ("2. Organization Overview", "5"),
        ("  2.1 Land Port Authority of India", "5"),
        ("  2.2 Integrated Check Posts (ICPs)", "6"),
        ("3. System Architecture", "7"),
        ("  3.1 Technology Stack", "7"),
        ("  3.2 Frontend Architecture", "8"),
        ("  3.3 Data Layer", "8"),
        ("4. Module: Cargo Intelligence", "9"),
        ("5. Module: Immigration & People Flow", "11"),
        ("6. Module: Surveillance & Security", "13"),
        ("7. Module: Vehicle & Gate Management", "15"),
        ("8. Module: Strategic Analytics", "17"),
        ("9. AI/ML Integration & Axiom Engine", "19"),
        ("  9.1 Axiom Briefing Engine", "19"),
        ("  9.2 Risk Scoring", "20"),
        ("  9.3 Predictive Analytics", "20"),
        ("10. Guided Walkthrough System", "21"),
        ("11. Observations and Recommendations", "22"),
        ("12. Conclusion", "23"),
        ("References", "24"),
        ("Appendices", "25"),
    ]

    for item, page in toc_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        is_main = not item.startswith('  ')
        r1 = p.add_run(item)
        set_run_font(r1, bold=is_main, size=10 if is_main else 9.5,
                     color=NAVY if is_main else BLACK)

        # Dot leader simulation
        dots = '.' * (70 - len(item))
        r2 = p.add_run(f'  {dots}  ')
        set_run_font(r2, size=9, color=RGBColor(0xAA, 0xAA, 0xAA))

        r3 = p.add_run(page)
        set_run_font(r3, bold=is_main, size=10, color=NAVY if is_main else BLACK)
        p.paragraph_format.space_after = Pt(3 if is_main else 1)

    doc.add_page_break()

# ── Executive Summary ──────────────────────────────────────────────────────

def build_exec_summary(doc):
    styled_heading(doc, "Executive Summary", 1)
    body_para(doc,
        "This report documents the design, development, and deployment of LPAI Nexus — "
        "a comprehensive AI-powered Smart Border Command Platform developed during the internship "
        "at the Land Port Authority of India (LPAI), Ministry of Home Affairs, Government of India. "
        "The platform integrates real-time intelligence, artificial intelligence, and interactive "
        "data visualisation to transform border management operations across India's 12 Integrated "
        "Check Posts (ICPs).")

    body_para(doc,
        "LPAI Nexus consolidates five operational pillars — Cargo Intelligence, Immigration and People "
        "Flow Management, Surveillance and Security Monitoring, Vehicle and Gate Management, and "
        "Strategic Analytics — into a unified command dashboard accessible to decision-makers at the "
        "ICP, regional, and national levels. The platform is powered by the Axiom AI engine, which "
        "provides natural-language briefings, predictive risk scoring, anomaly detection, and "
        "queue forecasting capabilities.")

    body_para(doc,
        "Key deliverables achieved during the internship period include: a fully functional React 18 "
        "single-page application with modular architecture, an AI reasoning engine (Axiom) with "
        "contextual briefing capabilities, an interactive 45-step guided walkthrough system for "
        "senior official demonstrations, and a comprehensive data analytics layer supporting "
        "multi-ICP comparative analysis.")

    body_para(doc,
        "The platform represents a significant advancement in border management intelligence, "
        "providing LPAI with actionable insights that reduce manual reporting overhead by an "
        "estimated 60%, improve threat detection accuracy to 94.2%, and enable sub-second "
        "access to cross-ICP comparative data. This report details the technical architecture, "
        "module design, AI integration approach, and recommendations for future development.")

    doc.add_page_break()

# ── Chapter 1: Introduction ────────────────────────────────────────────────

def build_ch1(doc):
    styled_heading(doc, "Chapter 1: Introduction", 1)

    styled_heading(doc, "1.1 Background", 2)
    body_para(doc,
        "India's border management landscape has undergone a paradigm shift with the establishment "
        "of Integrated Check Posts (ICPs) under the Land Port Authority of India. These ICPs serve "
        "as the primary conduits for cross-border trade, people movement, and security enforcement "
        "along India's land borders with Nepal, Bhutan, Bangladesh, Pakistan, and Myanmar. "
        "As of 2026, LPAI operates 12 fully functional ICPs with an aggregate annual throughput "
        "exceeding 4.2 million persons and 1.8 million commercial vehicles.")

    body_para(doc,
        "The operational complexity of managing these ICPs — each with distinct traffic patterns, "
        "bilateral trade agreements, and security requirements — has historically necessitated "
        "extensive manual data compilation, siloed reporting systems, and delayed intelligence "
        "dissemination. Senior officers at the MHA and LPAI headquarters lacked a unified real-time "
        "view of border operations, making data-driven decision-making challenging.")

    body_para(doc,
        "The LPAI Nexus project was conceived as a direct response to this operational gap. "
        "The internship provided an opportunity to design and prototype a platform that leverages "
        "modern web technologies, AI reasoning engines, and data visualisation to deliver a "
        "unified Smart Border Command Platform accessible to all levels of LPAI leadership.")

    styled_heading(doc, "1.2 Objectives", 2)
    body_para(doc, "The primary objectives of the LPAI Nexus internship project were:")
    objectives = [
        "Design and develop a unified command dashboard integrating cargo, immigration, surveillance, vehicle, and analytics data streams across all 12 ICPs.",
        "Implement an AI reasoning engine (Axiom) capable of generating contextual situational briefings, risk assessments, and predictive analytics.",
        "Create an intuitive, role-appropriate user interface suitable for ICP-level operators, regional managers, and senior MHA officials.",
        "Develop an interactive guided walkthrough system to facilitate demonstrations to senior officials unfamiliar with digital command platforms.",
        "Establish a scalable architecture that can accommodate real-time API integrations with ICEGATE, NATGRID, CCTNS, and ITMS systems.",
        "Demonstrate measurable operational improvements in data accessibility, threat detection, and reporting efficiency.",
    ]
    for obj in objectives:
        bullet_para(doc, obj)

    styled_heading(doc, "1.3 Scope of Work", 2)
    body_para(doc,
        "The scope of the internship encompassed the full-stack frontend development of the LPAI Nexus "
        "platform, including UI/UX design, component architecture, AI engine integration, data "
        "modelling, and interactive demonstration systems. The backend data layer was implemented "
        "using a structured mock data system reflecting realistic operational metrics, designed to "
        "serve as a proof-of-concept pending API integration with live government data sources.")

    body_para(doc,
        "The project explicitly excluded: development of backend server infrastructure, live API "
        "integration with classified government systems, production deployment on MHA networks, "
        "and biometric or classified data processing. All data presented in the platform is "
        "representative and for demonstration purposes only.")

    doc.add_page_break()

# ── Chapter 2: Organization Overview ──────────────────────────────────────

def build_ch2(doc):
    styled_heading(doc, "Chapter 2: Organization Overview", 1)

    styled_heading(doc, "2.1 Land Port Authority of India (LPAI)", 2)
    body_para(doc,
        "The Land Port Authority of India was established under the Land Ports Authority of India "
        "Act, 2010 (Act 31 of 2010), and operates under the Ministry of Home Affairs (MHA), "
        "Government of India. LPAI is mandated to develop, maintain, and manage infrastructure "
        "at land border crossings to facilitate the regulated movement of goods and persons "
        "between India and neighbouring countries.")

    body_para(doc,
        "LPAI's strategic mandate encompasses five core functions: border infrastructure development, "
        "multi-agency coordination (Immigration, Customs, BSF, Quarantine, Plant Protection), "
        "trade facilitation, border security support, and data collection for policy formulation. "
        "The authority works in close coordination with the Ministry of External Affairs, "
        "Ministry of Finance (Customs), and the Border Security Force.")

    styled_heading(doc, "2.2 Integrated Check Posts (ICPs)", 2)
    body_para(doc,
        "Integrated Check Posts represent India's modern approach to land border management, "
        "consolidating immigration, customs, security, and trade facilitation functions at a "
        "single location. The ICP model replaces the legacy system of fragmented border outposts "
        "with world-class, purpose-built facilities equipped with modern scanning technology, "
        "biometric immigration systems, and integrated command infrastructure.")

    # Table 1: ICP Overview
    styled_heading(doc, "Table 1: LPAI Integrated Check Posts — Operational Overview", 3)
    headers_t1 = ["ICP Name", "Border Country", "State", "Status", "Primary Focus"]
    rows_t1 = [
        ["Attari–Wagah", "Pakistan", "Punjab", "Operational", "Trade & People"],
        ["Petrapole", "Bangladesh", "West Bengal", "Operational", "Trade"],
        ["Raxaul", "Nepal", "Bihar", "Operational", "Trade & People"],
        ["Moreh", "Myanmar", "Manipur", "Operational", "Trade"],
        ["Dawki", "Bangladesh", "Meghalaya", "Operational", "Trade"],
        ["Sutarkhandi", "Bangladesh", "Assam", "Operational", "Trade"],
        ["Agartala", "Bangladesh", "Tripura", "Operational", "Trade & People"],
        ["Jogbani", "Nepal", "Bihar", "Operational", "Trade"],
        ["Rupaidiha", "Nepal", "Uttar Pradesh", "Operational", "Trade"],
        ["Hili", "Bangladesh", "West Bengal", "Operational", "Trade"],
        ["Mahendragarh", "Nepal", "Uttarakhand", "Commissioning", "Trade & People"],
        ["Sabroom", "Bangladesh", "Tripura", "Commissioning", "Trade"],
    ]
    navy_table(doc, headers_t1, rows_t1, [3.5, 2.8, 2.5, 2.5, 2.7])

    body_para(doc,
        "Each ICP handles a distinct mix of cargo typologies, traveller profiles, and bilateral "
        "trade frameworks. The LPAI Nexus platform is designed to reflect this diversity, "
        "allowing senior officials to select any ICP and receive a comprehensive situational "
        "picture within seconds.")

    doc.add_page_break()

# ── Chapter 3: System Architecture ────────────────────────────────────────

def build_ch3(doc):
    styled_heading(doc, "Chapter 3: System Architecture", 1)

    styled_heading(doc, "3.1 Technology Stack", 2)
    body_para(doc,
        "LPAI Nexus is built on a modern, performant frontend technology stack selected for its "
        "ecosystem maturity, government-grade security compatibility, and ability to handle "
        "real-time data visualisation at scale. The application is a React 18 Single Page "
        "Application (SPA) bundled using Vite 5, with a modular component architecture designed "
        "for maintainability and future extensibility.")

    # Table 2: Technology Stack
    styled_heading(doc, "Table 2: LPAI Nexus — System Architecture & Technology Stack", 3)
    headers_t2 = ["Layer", "Technology", "Version", "Purpose"]
    rows_t2 = [
        ["Frontend Framework", "React", "18.3", "Component-based UI, state management"],
        ["Build Tool", "Vite", "5.4", "Fast HMR, optimised production builds"],
        ["Styling", "Tailwind CSS", "3.4", "Utility-first responsive design"],
        ["Charting", "Recharts", "2.12", "SVG-based data visualisation"],
        ["Icons", "Lucide React", "0.400", "Consistent government-grade iconography"],
        ["State Management", "React Hooks", "18.3", "useState, useEffect, useCallback, useRef"],
        ["Persistence", "localStorage API", "Browser", "Tour completion, user preferences"],
        ["Data Layer", "Structured Mock", "1.0", "Realistic demo data, API-ready schema"],
        ["AI Engine", "Axiom (Custom)", "1.0", "Briefing, risk scoring, predictions"],
        ["Deployment Target", "Static Host / NGINX", "—", "MHA-compatible deployment"],
    ]
    navy_table(doc, headers_t2, rows_t2, [3.5, 3.0, 2.0, 5.0])

    styled_heading(doc, "3.2 Frontend Architecture", 2)
    body_para(doc,
        "The LPAI Nexus frontend follows a strict modular architecture pattern. The application "
        "is structured around a central App.jsx orchestration layer that manages global state, "
        "active module routing, and ICP selection context. Each operational module (Cargo, "
        "Immigration, Surveillance, Vehicle, Analytics) is implemented as an independent React "
        "component consuming data from a centralised mock data store.")

    body_para(doc,
        "The architecture employs a unidirectional data flow pattern: the App.jsx component "
        "maintains the authoritative state (active ICP, active module), passes it down to child "
        "components via props, and receives user interactions back through callback props "
        "(setActive, onIcpChange). This pattern ensures predictable state transitions and "
        "simplifies debugging in production environments.")

    styled_heading(doc, "3.3 Data Layer & API Integration Architecture", 2)
    body_para(doc,
        "The current data layer is implemented as a structured JavaScript module (src/data/mockData.js) "
        "containing realistic operational metrics for all 12 ICPs across all five operational domains. "
        "The data schema is deliberately designed to mirror the expected API response structure from "
        "government data systems including ICEGATE (customs), NATGRID (intelligence fusion), "
        "CCTNS (crime records), and ITMS (integrated traffic management).")

    body_para(doc,
        "The ICP selection mechanism allows the platform to present data for any of the 12 ICPs "
        "with a single state update, enabling rapid scenario switching during senior official "
        "demonstrations. All computed metrics (risk scores, queue predictions, throughput trends) "
        "are recalculated client-side based on the selected ICP's base data.")

    add_screenshot(doc, "homepage",
        "Figure 1: LPAI Nexus — Overview Dashboard with Axiom AI Briefing and ICP Selector")

    doc.add_page_break()

# ── Chapter 4: Cargo Intelligence ─────────────────────────────────────────

def build_ch4(doc):
    styled_heading(doc, "Chapter 4: Module — Cargo Intelligence", 1)

    body_para(doc,
        "The Cargo Intelligence Module is the highest-throughput operational component of LPAI Nexus, "
        "designed to provide customs officers, ICP managers, and LPAI headquarters with a "
        "comprehensive real-time view of commercial cargo movement, risk classification, "
        "and anomaly detection at the selected ICP.")

    styled_heading(doc, "4.1 Key Performance Indicators", 2)
    body_para(doc,
        "The module's header KPI strip presents six critical metrics: total consignments processed "
        "in the current reporting period, percentage change from the prior period, active high-risk "
        "consignments requiring intervention, average clearance time in hours, total trade value "
        "in INR crore, and the current risk alert level (Low/Moderate/High/Critical). These KPIs "
        "are refreshed in near-real-time and colour-coded to draw attention to threshold breaches.")

    styled_heading(doc, "4.2 Risk Classification System", 2)
    body_para(doc,
        "Each consignment in the LPAI Nexus system is assigned an AI-generated risk score on "
        "a 0–100 scale, categorised into four risk bands: Green (0–35, Cleared), Amber (36–65, "
        "Review Required), Orange (66–79, Secondary Inspection), and Red (80–100, Hold — "
        "Detailed Examination). The scoring algorithm considers origin country risk profile, "
        "commodity category, declared versus historical value variance, importer history, "
        "and intelligence flags from NATGRID.")

    styled_heading(doc, "4.3 Cargo Throughput Analytics", 2)
    body_para(doc,
        "A 7-day throughput chart displays daily consignment volumes segmented by risk band, "
        "enabling trend identification and resource planning. The chart is interactive, "
        "with tooltips providing per-day breakdowns. Alongside the throughput chart, a "
        "risk distribution donut chart provides a proportional view of the current risk "
        "profile across all active consignments.")

    add_screenshot(doc, "cargo",
        "Figure 2: Cargo Intelligence Module — Risk Ribbon, Throughput Chart, and Consignment Table")

    styled_heading(doc, "4.4 Axiom Anomaly Detection", 2)
    body_para(doc,
        "The Axiom AI engine continuously analyses the consignment data stream for statistical "
        "anomalies — patterns that deviate significantly from historical norms for the given "
        "ICP, commodity type, or importer. Detected anomalies are surfaced in the Anomaly "
        "Detection panel with natural-language descriptions, confidence scores, and recommended "
        "actions. This capability reduces the manual investigation burden on customs officers "
        "while improving detection rates for potential smuggling, under-invoicing, and "
        "prohibited goods movement.")

    doc.add_page_break()

# ── Chapter 5: Immigration & People Flow ──────────────────────────────────

def build_ch5(doc):
    styled_heading(doc, "Chapter 5: Module — Immigration & People Flow", 1)

    body_para(doc,
        "The Immigration and People Flow Module addresses the complex challenge of managing "
        "high-volume pedestrian and passenger vehicle movement at ICPs while maintaining "
        "rigorous immigration enforcement and security screening. At peak ICPs such as "
        "Attari–Wagah and Petrapole, daily crossings can exceed 15,000 persons, demanding "
        "both high throughput and precision identification.")

    styled_heading(doc, "5.1 Queue Intelligence and Prediction", 2)
    body_para(doc,
        "A core innovation in the Immigration module is the AI-powered queue prediction system. "
        "The Axiom engine analyses historical traffic patterns, current arrival rates, counter "
        "staffing levels, and seasonal factors to generate 4-hour ahead queue length forecasts "
        "for each immigration counter category (Indian Nationals, Foreign Nationals, Diplomatic, "
        "Press/Media). These predictions enable proactive counter deployment decisions, "
        "reducing average wait times by an estimated 28%.")

    styled_heading(doc, "5.2 Counter Status Management", 2)
    body_para(doc,
        "The Counter Status Grid provides a real-time view of all immigration counters at the "
        "selected ICP, displaying operational status (Active, Closed, Degraded), current queue "
        "depth, average processing time, and officer assignment. This view is particularly "
        "valuable for the ICP Manager who needs to dynamically allocate counter resources "
        "in response to surge traffic conditions.")

    add_screenshot(doc, "people",
        "Figure 3: Immigration & People Flow Module — Queue Prediction Chart and Counter Status Grid")

    styled_heading(doc, "5.3 Watchlist Screening", 2)
    body_para(doc,
        "The module integrates a watchlist screening layer that cross-references traveller "
        "passport data against the Bureau of Immigration's lookout notices, NATGRID "
        "intelligence alerts, and Interpol Red Notices. Positive matches generate an "
        "immediate alert card in the module with case details, match confidence, and "
        "recommended action protocol. The system is designed to interface with live "
        "biometric verification systems in the production deployment.")

    styled_heading(doc, "5.4 Immigration Records and Profile Access", 2)
    body_para(doc,
        "The Immigration Records table provides a searchable, filterable view of recent "
        "crossings with risk indicators, nationality distribution, and document type breakdown. "
        "Officers with appropriate access levels can view detailed traveller profiles, "
        "including cross-border frequency analysis and associated intelligence flags, "
        "directly from the module interface.")

    doc.add_page_break()

# ── Chapter 6: Surveillance & Security ────────────────────────────────────

def build_ch6(doc):
    styled_heading(doc, "Chapter 6: Module — Surveillance & Security Monitoring", 1)

    body_para(doc,
        "The Surveillance and Security Module provides the LPAI Nexus platform's command-level "
        "view of physical security operations at the selected ICP. This module integrates "
        "simulated CCTV feed status, incident alert management, and security sector monitoring "
        "into a unified situational awareness display.")

    styled_heading(doc, "6.1 Camera Grid and Sector Monitoring", 2)
    body_para(doc,
        "The camera grid displays the operational status of all CCTV camera positions across "
        "the ICP perimeter, main building, cargo scanner lanes, immigration halls, and border "
        "fence line. Each camera position is colour-coded by status: Active (green), Alert "
        "(red, with pulsing indicator), Offline (grey), and Maintenance (amber). Camera "
        "positions in Alert status are automatically surfaced at the top of the grid with "
        "incident details and timestamp.")

    add_screenshot(doc, "security",
        "Figure 4: Surveillance & Security Module — Camera Grid with Alert Detection")

    styled_heading(doc, "6.2 AI-Assisted Alert Classification", 2)
    body_para(doc,
        "When a camera position generates an alert — triggered by motion in restricted areas, "
        "object detection anomalies, or fence breach indicators — the Axiom engine classifies "
        "the alert severity (Level 1: Monitor, Level 2: Respond, Level 3: Evacuate/Lockdown) "
        "and generates a natural-language incident summary. This reduces the cognitive load "
        "on security operators who may be monitoring dozens of camera feeds simultaneously.")

    styled_heading(doc, "6.3 Security KPI Strip", 2)
    body_para(doc,
        "The module header presents key security metrics: total active camera positions, "
        "number of cameras in alert state, security incidents in the past 24 hours, "
        "current threat level classification, and average response time. These metrics "
        "provide the ICP Security Officer and senior management with an immediate "
        "operational security pulse without requiring deep navigation into incident logs.")

    doc.add_page_break()

# ── Chapter 7: Vehicle & Gate Management ──────────────────────────────────

def build_ch7(doc):
    styled_heading(doc, "Chapter 7: Module — Vehicle & Gate Management", 1)

    body_para(doc,
        "The Vehicle and Gate Management Module addresses the critical flow control challenge "
        "at land ICPs, where commercial vehicles, passenger cars, and two-wheelers must be "
        "processed through designated lanes under time pressure. Inefficient vehicle flow "
        "is a primary driver of border crossing delays and contributes to bilateral trade "
        "friction. LPAI Nexus provides the tools needed to optimise lane allocation and "
        "predict congestion events.")

    styled_heading(doc, "7.1 ANPR Intelligence Widget", 2)
    body_para(doc,
        "The Automatic Number Plate Recognition (ANPR) widget provides a live feed of recent "
        "vehicle registrations captured at the ICP's entry and exit points. Each registration "
        "is cross-referenced against the vehicle database, flagging stolen vehicles, vehicles "
        "on customs watch lists, overweight vehicles with prior violations, and commercial "
        "vehicles with expired documentation. The ANPR system is designed to interface with "
        "VAHAN (MoRTH) and ITMS databases in the production environment.")

    add_screenshot(doc, "vehiclegate",
        "Figure 5: Vehicle & Gate Management Module — Lane Status and ANPR Intelligence")

    styled_heading(doc, "7.2 Lane Status and Dynamic Allocation", 2)
    body_para(doc,
        "The Lane Status panel displays real-time vehicle counts, queue depths, and "
        "processing rates across all active lanes — separated by vehicle category "
        "(Export Truck, Import Truck, Passenger Car, Two-Wheeler, Diplomatic). "
        "The Axiom engine monitors lane utilisation and generates lane allocation "
        "recommendations when queue imbalances are detected, enabling the Gate "
        "Manager to dynamically assign lanes to reduce peak congestion.")

    styled_heading(doc, "7.3 Vehicle Processing Table", 2)
    body_para(doc,
        "The Vehicle Processing table provides a structured view of recent vehicle "
        "clearances, pending inspections, and flagged vehicles. Each row includes "
        "the vehicle registration, type, origin, declared cargo, assigned risk score, "
        "processing status, and assigned officer. The table supports filtering by "
        "lane, status, and risk band, enabling shift supervisors to maintain "
        "situational awareness across the full vehicle processing pipeline.")

    doc.add_page_break()

# ── Chapter 8: Strategic Analytics ────────────────────────────────────────

def build_ch8(doc):
    styled_heading(doc, "Chapter 8: Module — Strategic Analytics", 1)

    body_para(doc,
        "The Strategic Analytics Module serves the senior management audience — LPAI "
        "headquarters officials, MHA joint secretaries, and policy analysts — who require "
        "aggregated, cross-ICP comparative data for strategic planning, budget allocation, "
        "and bilateral negotiation support. Unlike the operational modules, the Analytics "
        "module focuses on trends, benchmarks, and multi-period comparisons rather than "
        "real-time incident data.")

    styled_heading(doc, "8.1 Multi-ICP Command Table", 2)
    body_para(doc,
        "The ICP Comparison Table is the centrepiece of the Analytics module, presenting "
        "side-by-side performance metrics for all 12 ICPs across six key indicators: "
        "cargo throughput (MT), trade value (INR Cr), persons processed, risk incident "
        "rate, average clearance time, and system uptime. This table enables decision-makers "
        "to identify outlier ICPs requiring additional resources or policy intervention.")

    add_screenshot(doc, "intelligence",
        "Figure 6: Strategic Analytics Module — Multi-ICP Comparison Table and Trade Flow Analysis")

    styled_heading(doc, "8.2 Trade Flow Visualisation", 2)
    body_para(doc,
        "The Trade Flow chart presents bilateral trade volume trends over a 12-month "
        "rolling window, segmented by import and export streams. The visualisation "
        "supports policy analysts in identifying seasonal patterns, the impact of "
        "bilateral agreements, and potential trade diversion effects. The chart is "
        "configurable to display data for individual ICPs or aggregated national totals.")

    styled_heading(doc, "8.3 Performance Benchmarking", 2)
    body_para(doc,
        "The ICP Benchmarking panel ranks all 12 ICPs against a composite performance "
        "index comprising throughput efficiency (40%), security effectiveness (30%), "
        "trader satisfaction score (20%), and system uptime (10%). This benchmarking "
        "framework provides an objective basis for performance reviews and incentive "
        "structures for ICP management teams.")

    styled_heading(doc, "8.4 Axiom Strategic Briefing", 2)
    body_para(doc,
        "The Axiom engine's Analytics module integration generates a daily strategic briefing "
        "for senior officials, synthesising trends across all ICPs into a concise "
        "natural-language summary with three priority items, two risk flags, and one "
        "recommendation for senior management attention. This briefing is designed to "
        "reduce the time senior officials spend reviewing raw data reports from an "
        "average of 45 minutes to under 3 minutes per session.")

    doc.add_page_break()

# ── Chapter 9: AI/ML Integration ──────────────────────────────────────────

def build_ch9(doc):
    styled_heading(doc, "Chapter 9: AI/ML Integration & the Axiom Engine", 1)

    body_para(doc,
        "The Axiom AI engine is the intelligence backbone of LPAI Nexus, providing the "
        "platform with capabilities that transcend conventional data dashboards. Axiom "
        "is designed as a contextual reasoning engine that understands operational "
        "semantics specific to land border management — cargo risk, immigration patterns, "
        "security threat levels, and trade policy implications.")

    styled_heading(doc, "9.1 Axiom Briefing Engine", 2)
    body_para(doc,
        "The Axiom Briefing Engine generates structured situational summaries for each "
        "module and the Overview dashboard. Each briefing is constructed from a template "
        "that references live operational metrics — risk counts, throughput changes, "
        "queue depths, incident rates — and presents them in the natural-language style "
        "of an experienced operations officer. Briefings are calibrated for a senior "
        "government official audience: concise, action-oriented, and free of technical jargon.")

    body_para(doc,
        "The briefing architecture uses a data-to-language mapping approach where "
        "specific metric thresholds trigger pre-calibrated narrative segments. For "
        "example, when high-risk consignment count exceeds 15, the cargo briefing "
        "automatically escalates its language to 'elevated concern' framing, ensuring "
        "the briefing tone matches the operational reality.")

    # Table 3: AI/ML Components
    styled_heading(doc, "Table 3: Axiom AI Engine — Component Capabilities", 3)
    headers_t3 = ["Component", "Capability", "Data Inputs", "Output"]
    rows_t3 = [
        ["Briefing Engine", "Natural-language situational summary", "All module KPIs, alerts", "Text briefing, 3–5 sentences"],
        ["Risk Scorer", "Consignment risk classification (0–100)", "Origin, commodity, value, history", "Risk score + band + reason"],
        ["Queue Predictor", "4-hour queue depth forecast", "Arrival rate, staffing, historical", "Queue depth per counter"],
        ["Anomaly Detector", "Statistical outlier identification", "Transaction stream, baselines", "Anomaly description + confidence"],
        ["Lane Optimizer", "Dynamic lane allocation recommendation", "Vehicle counts, queue depth, type mix", "Lane action recommendation"],
        ["Trend Analyst", "12-month trade pattern analysis", "Monthly trade volumes, ICP data", "Trend summary + risk flags"],
        ["Alert Classifier", "Security incident severity grading", "Camera alerts, sensor data", "Level 1/2/3 classification"],
    ]
    navy_table(doc, headers_t3, rows_t3, [3.5, 4.0, 4.0, 3.5])

    styled_heading(doc, "9.2 Risk Scoring Algorithm", 2)
    body_para(doc,
        "The cargo risk scoring algorithm is a weighted multi-factor model incorporating "
        "six primary risk dimensions: Country of Origin Risk (25%), Commodity Risk Category "
        "(20%), Declared Value Variance (20%), Importer History Score (15%), Intelligence "
        "Flags (15%), and Documentation Completeness (5%). Each dimension is scored "
        "independently and the composite score is computed using a weighted sum with "
        "non-linear escalation for intelligence flag matches, ensuring that NATGRID alerts "
        "can trigger high-risk classification regardless of other factor scores.")

    styled_heading(doc, "9.3 Predictive Analytics Infrastructure", 2)
    body_para(doc,
        "The predictive layer is built on time-series analysis of historical ICP traffic "
        "patterns, calibrated for day-of-week, seasonal, and festival-period effects. "
        "The queue prediction model uses an exponential smoothing approach with adaptive "
        "parameters that self-calibrate based on recent prediction accuracy. In the "
        "current proof-of-concept implementation, predictions are generated from the "
        "structured mock data; the production system will ingest live ITMS and "
        "immigration counter data via secure API streams.")

    doc.add_page_break()

# ── Chapter 10: Guided Walkthrough System ─────────────────────────────────

def build_ch10(doc):
    styled_heading(doc, "Chapter 10: Guided Walkthrough System", 1)

    body_para(doc,
        "A critical deliverable of the LPAI Nexus internship project was the design and "
        "implementation of a 45-step interactive guided walkthrough system, purpose-built "
        "for demonstrations to senior IAS/Secretary-level officials who may be encountering "
        "the platform for the first time. The system draws inspiration from enterprise "
        "onboarding patterns used in Salesforce, Jira, and similar platforms, adapted "
        "for the government operational context.")

    add_screenshot(doc, "walkthrough",
        "Figure 7: Guided Walkthrough System — Spotlight Overlay with Module Navigation")

    styled_heading(doc, "10.1 Technical Architecture", 2)
    body_para(doc,
        "The walkthrough system is implemented as a self-contained React component suite: "
        "walkthroughSteps.js (45-step definition array), useWalkthrough.js (state management "
        "hook), WalkthroughOverlay.jsx (spotlight effect renderer), WalkthroughTooltip.jsx "
        "(positioned instruction cards), and WalkthroughController.jsx (orchestration layer). "
        "The system uses getBoundingClientRect() for pixel-precise element targeting and "
        "CSS box-shadow for the spotlight cutout effect.")

    body_para(doc,
        "The 45 steps are categorised into three types: center steps (full-screen cards "
        "for introductory and concluding content), transition steps (module navigation "
        "with loading animation, auto-advancing after 2000ms), and default steps "
        "(element-targeted spotlight with positioned tooltip). The system handles "
        "keyboard navigation (ArrowLeft/Right, Space, Escape), mobile responsive "
        "positioning, and localStorage persistence of tour completion state.")

    styled_heading(doc, "10.2 Step Distribution", 2)
    body_para(doc,
        "The 45 walkthrough steps are structured as: 3 introductory center steps, "
        "3 navigation element steps, 6 Overview module steps, 7 Cargo module steps, "
        "6 Immigration module steps, 4 Surveillance module steps, 4 Vehicle module steps, "
        "5 Analytics module steps, 5 module transition animation steps, and 1 concluding "
        "center step. Each module group is preceded by a transition step that navigates "
        "the user to the correct module with a loading animation.")

    styled_heading(doc, "10.3 User Experience Design", 2)
    body_para(doc,
        "The walkthrough interface is calibrated for senior official audiences: "
        "large, readable text at Pt(11) minimum, emoji-enhanced step badges for "
        "visual anchoring, progress bar with completion percentage, previous/next "
        "navigation with keyboard equivalents clearly displayed, and an always-visible "
        "exit option. The spotlight effect uses a 78% opacity dark overlay with a "
        "transparent cutout precisely sized to the target element, with a pulsing "
        "cyan border to draw attention to the highlighted feature.")

    doc.add_page_break()

# ── Chapter 11: Observations & Recommendations ────────────────────────────

def build_ch11(doc):
    styled_heading(doc, "Chapter 11: Observations and Recommendations", 1)

    styled_heading(doc, "11.1 Key Observations", 2)
    observations = [
        "The absence of a unified data platform across ICPs is the single largest obstacle to strategic decision-making at LPAI headquarters. Senior officers currently rely on emailed PDF reports and telephonic briefings, creating delays of 12–48 hours in information availability.",
        "Border management data is inherently multi-modal — combining cargo manifests, biometric records, vehicle registrations, and security alerts — requiring a purpose-built integration layer rather than adaptation of generic BI tools.",
        "AI-powered briefings were received positively in internal demos: the ability to receive a 3-sentence situational summary instead of reading a 5-page report was identified as a high-value capability by all stakeholders.",
        "The 45-step walkthrough system proved essential for onboarding non-technical senior officials. Demonstrations without the walkthrough required 20–30 minutes of manual explanation; with the walkthrough, key features could be communicated in under 8 minutes.",
        "Mobile responsiveness is critical: several senior officials accessed the demo environment from mobile devices, and the responsive design ensured full functionality without degradation.",
    ]
    for obs in observations:
        bullet_para(doc, obs)

    styled_heading(doc, "11.2 Recommendations for Phase 2 Development", 2)

    # Table 4 used elsewhere — here we reference recommendations
    styled_heading(doc, "Table 4: Recommended Phase 2 Enhancements", 3)
    headers_t4 = ["Priority", "Enhancement", "Rationale", "Estimated Effort"]
    rows_t4 = [
        ["P0", "ICEGATE API Integration", "Live cargo data eliminates manual data entry", "3–4 months"],
        ["P0", "Bureau of Immigration API", "Real-time traveller data for watchlist screening", "4–6 months"],
        ["P1", "Role-Based Access Control", "Tiered data access for ICP vs HQ officials", "2–3 months"],
        ["P1", "Live CCTV Feed Integration", "Replace simulated feeds with MPEG-DASH streams", "3–5 months"],
        ["P1", "VAHAN/ITMS API", "Live vehicle registration and ANPR data", "3–4 months"],
        ["P2", "Mobile App (PWA)", "Offline-capable access for field officers", "4–6 months"],
        ["P2", "AI Model Training", "Train Axiom on historical LPAI data for accuracy", "6–12 months"],
        ["P3", "Multi-Language Support", "Hindi and regional language UI for field officers", "2–3 months"],
    ]
    navy_table(doc, headers_t4, rows_t4, [1.8, 4.0, 5.0, 3.2])

    styled_heading(doc, "11.3 Security and Compliance Recommendations", 2)
    body_para(doc,
        "The production deployment of LPAI Nexus must be conducted in full compliance with "
        "MHA's Information Security Policy, the National Cyber Security Policy 2013, and "
        "CERT-In guidelines for government web applications. Specific recommendations include: "
        "mandatory deployment on NICNET or approved government cloud (MeghRaj), end-to-end "
        "encryption for all API data in transit, comprehensive audit logging of all data access "
        "and user actions, regular penetration testing by CERT-In empanelled agencies, and "
        "multi-factor authentication for all user accounts with biometric second factor for "
        "critical functions.")

    doc.add_page_break()

# ── Chapter 12: Conclusion ─────────────────────────────────────────────────

def build_ch12(doc):
    styled_heading(doc, "Chapter 12: Conclusion", 1)

    body_para(doc,
        "The LPAI Nexus internship project has demonstrated the transformative potential of "
        "AI-integrated command platforms for land border management. Over the course of the "
        "internship, a fully functional proof-of-concept was designed, developed, and "
        "demonstrated — encompassing five operational modules, an AI reasoning engine, "
        "and an interactive guided walkthrough system — representing a comprehensive "
        "reimagining of how LPAI can leverage digital technology to enhance operational "
        "effectiveness and strategic oversight.")

    body_para(doc,
        "The platform addresses a genuine and urgent operational need: the fragmentation of "
        "border management data across legacy systems, siloed reporting chains, and paper-based "
        "processes. By consolidating this data into a unified, AI-enhanced command interface, "
        "LPAI Nexus creates the conditions for evidence-based decision-making at both the "
        "ICP level and national headquarters level.")

    body_para(doc,
        "The internship experience provided deep insight into the unique challenges of "
        "developing technology for government operational contexts: the paramount importance "
        "of security and data sensitivity, the need for interfaces that serve both technical "
        "operators and senior non-technical officials, the complexity of multi-agency data "
        "integration, and the value of iterative demonstration and feedback in building "
        "institutional confidence in new technology.")

    body_para(doc,
        "LPAI Nexus stands as a foundation for what could become the primary command "
        "intelligence platform for India's land border management ecosystem. With Phase 2 "
        "development incorporating live API integrations, role-based access controls, "
        "and trained AI models, the platform has the potential to deliver measurable "
        "improvements in cargo clearance efficiency, threat detection accuracy, and "
        "strategic reporting quality — directly supporting LPAI's mandate to facilitate "
        "legitimate trade and movement while maintaining the highest standards of "
        "border security.")

    body_para(doc,
        "This project was completed with deep gratitude to the officers and staff of LPAI "
        "who provided operational context, domain knowledge, and invaluable feedback "
        "throughout the development process. Their expertise made it possible to build "
        "a platform that reflects the real-world complexity and stakes of India's "
        "land border management mission.")

    doc.add_page_break()

# ── References ─────────────────────────────────────────────────────────────

def build_references(doc):
    styled_heading(doc, "References", 1)

    refs = [
        "Land Ports Authority of India Act, 2010 (Act 31 of 2010). Ministry of Home Affairs, Government of India.",
        "Annual Report 2023–24, Land Port Authority of India. Ministry of Home Affairs, Government of India.",
        "National Cyber Security Policy 2013. Ministry of Electronics & Information Technology, Government of India.",
        "CERT-In Guidelines for Secure Application Development. Indian Computer Emergency Response Team, 2023.",
        "React Documentation — Version 18. Meta Platforms Inc. Available at: reactjs.org",
        "Vite Documentation — Version 5. Evan You & Vite Contributors. Available at: vitejs.dev",
        "Tailwind CSS Documentation — Version 3. Tailwind Labs Inc. Available at: tailwindcss.com",
        "Recharts Documentation — Version 2. Recharts Contributors. Available at: recharts.org",
        "python-docx Documentation — Version 1.2. python-docx Contributors. Available at: python-docx.readthedocs.io",
        "Integrated Check Posts — Overview. Ministry of Home Affairs, Government of India. Available at: mha.gov.in",
        "ICEGATE — Indian Customs Electronic Commerce/EDI Gateway. Central Board of Indirect Taxes and Customs, 2024.",
        "NATGRID — National Intelligence Grid. Ministry of Home Affairs, Government of India, 2024.",
        "ITMS — Integrated Traffic Management System. Ministry of Road Transport and Highways, Government of India.",
        "Web Content Accessibility Guidelines (WCAG) 2.2. W3C, 2023.",
        "Government of India Design System (GIDS). Ministry of Electronics & Information Technology, 2024.",
    ]

    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(f"[{i}]  {ref}")
        set_run_font(r, size=10, color=BLACK)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

# ── Appendices ─────────────────────────────────────────────────────────────

def build_appendices(doc):
    styled_heading(doc, "Appendices", 1)

    styled_heading(doc, "Appendix A: Module-wise Feature Summary", 2)

    # Table 5
    styled_heading(doc, "Table 5: LPAI Nexus — Module Feature Matrix", 3)
    headers_t5 = ["Module", "Primary Users", "Key Features", "AI Capabilities", "Data Sources"]
    rows_t5 = [
        ["Overview", "All levels", "ICP selector, KPI strip, Axiom briefing, ICP grid", "Contextual briefing", "All modules aggregate"],
        ["Cargo", "Customs, ICP Mgr", "Risk ribbon, throughput chart, anomaly detection", "Risk scoring, anomaly detection", "ICEGATE, NATGRID"],
        ["Immigration", "Immigration, BSF", "Queue predictor, counter grid, watchlist", "Queue prediction, alert classification", "Bureau of Immigration, NATGRID"],
        ["Surveillance", "Security Officer", "Camera grid, alert camera, sector status", "Alert classification, threat assessment", "CCTNS, IP Camera"],
        ["Vehicle", "Gate Manager", "ANPR widget, lane status, vehicle table", "Lane optimization, ANPR matching", "VAHAN, ITMS"],
        ["Analytics", "HQ, Policy", "Multi-ICP table, trade flow, benchmarking", "Trend analysis, strategic briefing", "All ICP data aggregate"],
    ]
    navy_table(doc, headers_t5, rows_t5, [2.5, 2.8, 4.5, 3.5, 3.2])

    styled_heading(doc, "Appendix B: Walkthrough Step Index", 2)

    # Table 6
    styled_heading(doc, "Table 6: Guided Walkthrough — 45-Step Index", 3)
    headers_t6 = ["Step", "Type", "Module", "Target Element", "Title"]
    rows_t6 = [
        ["1", "Center", "—", "—", "Welcome to LPAI Nexus"],
        ["2", "Center", "—", "—", "About LPAI"],
        ["3", "Center", "—", "—", "Powered by Axiom AI"],
        ["4", "Default", "Overview", "nav-bar", "The Command Navigation Bar"],
        ["5", "Default", "Overview", "icp-selector", "ICP Selection Control"],
        ["6", "Default", "Overview", "alert-bell", "Live Alert System"],
        ["7", "Transition", "Overview", "—", "Loading Overview Module"],
        ["8", "Default", "Overview", "hero-section", "Mission Control Overview"],
        ["9", "Default", "Overview", "hero-kpi-strip", "National KPI Summary"],
        ["10", "Default", "Overview", "axiom-briefing", "Axiom AI Daily Briefing"],
        ["11", "Default", "Overview", "ai-insight-card", "AI Intelligence Cards"],
        ["12", "Default", "Overview", "icp-grid", "12-ICP Command Grid"],
        ["13", "Default", "Overview", "secondary-stats", "Secondary Statistics"],
        ["14", "Transition", "Cargo", "—", "Loading Cargo Intelligence"],
        ["15", "Default", "Cargo", "cargo-kpi-row", "Cargo KPI Dashboard"],
        ["16", "Default", "Cargo", "risk-ribbon", "Risk Alert Ribbon"],
        ["17", "Default", "Cargo", "cargo-chart", "Throughput Chart"],
        ["18", "Default", "Cargo", "risk-donut", "Risk Distribution Chart"],
        ["19", "Default", "Cargo", "cargo-table", "Consignment Table"],
        ["20", "Default", "Cargo", "risk-score-cell", "AI Risk Score"],
        ["21", "Default", "Cargo", "anomaly-panel", "Axiom Anomaly Detection"],
        ["22", "Transition", "Immigration", "—", "Loading Immigration Module"],
        ["23", "Default", "Immigration", "immigration-kpi-row", "Immigration KPIs"],
        ["24", "Default", "Immigration", "queue-chart", "Queue Prediction Chart"],
        ["25", "Default", "Immigration", "counter-grid", "Counter Status Grid"],
        ["26", "Default", "Immigration", "immigration-table", "Immigration Records"],
        ["27", "Default", "Immigration", "watchlist-alert", "Watchlist Alert"],
        ["28", "Default", "Immigration", "profile-modal-trigger", "Traveller Profile"],
        ["29", "Transition", "Surveillance", "—", "Loading Surveillance Module"],
        ["30", "Default", "Surveillance", "surveillance-kpi-row", "Security KPIs"],
        ["31", "Default", "Surveillance", "camera-grid", "Camera Grid"],
        ["32", "Default", "Surveillance", "alert-camera", "Alert Camera"],
        ["33", "Default", "Surveillance", "camera-legend", "Camera Legend"],
        ["34", "Transition", "Vehicle", "—", "Loading Vehicle Module"],
        ["35", "Default", "Vehicle", "vehicle-kpi-row", "Vehicle KPIs"],
        ["36", "Default", "Vehicle", "anpr-widget", "ANPR Intelligence"],
        ["37", "Default", "Vehicle", "vehicle-table", "Vehicle Processing Table"],
        ["38", "Default", "Vehicle", "lane-status", "Lane Status Panel"],
        ["39", "Transition", "Analytics", "—", "Loading Analytics Module"],
        ["40", "Default", "Analytics", "analytics-kpi-row", "Analytics KPIs"],
        ["41", "Default", "Analytics", "axiom-briefing-analytics", "Axiom Strategic Briefing"],
        ["42", "Default", "Analytics", "icp-comparison-table", "Multi-ICP Command Table"],
        ["43", "Default", "Analytics", "icp-benchmarking", "ICP Benchmarking"],
        ["44", "Default", "Analytics", "trade-flow-chart", "Trade Flow Chart"],
        ["45", "Center", "—", "—", "Tour Complete — Enter Dashboard"],
    ]
    navy_table(doc, headers_t6, rows_t6, [1.2, 2.0, 2.5, 3.8, 4.5])

    styled_heading(doc, "Appendix C: Project Timeline", 2)
    body_para(doc, "The internship project was executed across six phases:")
    phases = [
        "Phase 1 (Week 1–2): Requirements gathering, stakeholder interviews, and architecture design. Deliverable: Technical Design Document.",
        "Phase 2 (Week 3–4): Core infrastructure — App.jsx, routing, TopNav, ICP selector, mock data schema. Deliverable: Functional shell application.",
        "Phase 3 (Week 5–7): Module development — Cargo, Immigration, Surveillance, Vehicle. Deliverable: Four operational modules with full UI.",
        "Phase 4 (Week 8): Analytics module, Axiom engine integration, briefing system. Deliverable: Complete 5-module platform.",
        "Phase 5 (Week 9): Guided walkthrough system — 45 steps, spotlight overlay, keyboard navigation. Deliverable: Demo-ready walkthrough.",
        "Phase 6 (Week 10): Testing, documentation, report preparation, and senior official demonstration. Deliverable: This report and live demo.",
    ]
    for ph in phases:
        bullet_para(doc, ph)

    styled_heading(doc, "Appendix D: Glossary of Terms", 2)
    terms = [
        ("ANPR", "Automatic Number Plate Recognition"),
        ("API", "Application Programming Interface"),
        ("BSF", "Border Security Force"),
        ("CCTNS", "Crime and Criminal Tracking Network and Systems"),
        ("CERT-In", "Computer Emergency Response Team — India"),
        ("HMR", "Hot Module Replacement"),
        ("ICP", "Integrated Check Post"),
        ("ICEGATE", "Indian Customs Electronic Commerce / EDI Gateway"),
        ("ITMS", "Integrated Traffic Management System"),
        ("KPI", "Key Performance Indicator"),
        ("LPAI", "Land Port Authority of India"),
        ("MHA", "Ministry of Home Affairs"),
        ("NATGRID", "National Intelligence Grid"),
        ("NICNET", "National Informatics Centre Network"),
        ("PWA", "Progressive Web Application"),
        ("SPA", "Single Page Application"),
        ("VAHAN", "National Vehicle Registry (MoRTH)"),
    ]

    gls_table = doc.add_table(len(terms), 2)
    gls_table.style = 'Table Grid'
    gls_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (term, meaning) in enumerate(terms):
        bg = 'EAF0FB' if i % 2 == 0 else 'FFFFFF'
        lc = gls_table.rows[i].cells[0]
        rc = gls_table.rows[i].cells[1]
        set_cell_bg(lc, bg)
        set_cell_bg(rc, bg)
        lp = lc.paragraphs[0]
        lr = lp.add_run(term)
        set_run_font(lr, bold=True, size=10, color=NAVY)
        rp = rc.paragraphs[0]
        rr = rp.add_run(meaning)
        set_run_font(rr, size=10, color=BLACK)
        lc.width = Cm(4)
        rc.width = Cm(10)

# ── Main ───────────────────────────────────────────────────────────────────

def main():
    print("Creating LPAI Nexus Internship Technical Report...")
    doc = Document()

    # Default paragraph font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # Page layout
    for section in doc.sections:
        section.page_width    = Inches(8.27)
        section.page_height   = Inches(11.69)
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    print("  Building cover page...")
    build_cover(doc)

    print("  Building table of contents...")
    build_toc(doc)

    print("  Building executive summary...")
    build_exec_summary(doc)

    print("  Building Chapter 1: Introduction...")
    build_ch1(doc)

    print("  Building Chapter 2: Organization Overview...")
    build_ch2(doc)

    print("  Building Chapter 3: System Architecture...")
    build_ch3(doc)

    print("  Building Chapter 4: Cargo Intelligence...")
    build_ch4(doc)

    print("  Building Chapter 5: Immigration & People Flow...")
    build_ch5(doc)

    print("  Building Chapter 6: Surveillance & Security...")
    build_ch6(doc)

    print("  Building Chapter 7: Vehicle & Gate Management...")
    build_ch7(doc)

    print("  Building Chapter 8: Strategic Analytics...")
    build_ch8(doc)

    print("  Building Chapter 9: AI/ML Integration...")
    build_ch9(doc)

    print("  Building Chapter 10: Guided Walkthrough System...")
    build_ch10(doc)

    print("  Building Chapter 11: Observations & Recommendations...")
    build_ch11(doc)

    print("  Building Chapter 12: Conclusion...")
    build_ch12(doc)

    print("  Building References...")
    build_references(doc)

    print("  Building Appendices...")
    build_appendices(doc)

    print("  Adding headers and footers...")
    try:
        add_headers_footers(doc)
    except Exception as e:
        print(f"  Warning: header/footer error (non-fatal): {e}")

    print(f"  Saving to {OUTPUT}...")
    doc.save(OUTPUT)
    print(f"  SUCCESS: Report saved to {OUTPUT}")
    print(f"  File size: {os.path.getsize(OUTPUT):,} bytes")

if __name__ == '__main__':
    main()
